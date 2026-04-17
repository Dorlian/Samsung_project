# -*- coding: utf-8 -*-
"""
Генерация Word-отчёта: датасет MRL, сравнение ResNet-18 / MobileNetV3, мини-выводы к графикам, пример реализации.
Запуск из корня репозитория: python scripts/build_eye_cnn_report_docx.py
Требует: pip install python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = cell


def _add_picture_safe(doc: Document, path: Path, caption: str) -> None:
    if not path.is_file():
        p = doc.add_paragraph()
        r = p.add_run(f"(Файл не найден: {path})")
        r.italic = True
        return
    doc.add_picture(str(path), width=Inches(5.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    assets = root / "assets"
    out = root / "Отчёт_нейросети_глаза_и_данные.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    t = doc.add_paragraph("Фотоассистент: нейросетевые модели глаз и данные для обучения")
    t.runs[0].bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Документ формируется автоматически; графики берутся из каталога assets/ при наличии файлов."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Сведения о данных (датасет и модели в проекте)", level=1)
    doc.add_paragraph(
        "Ниже — сводная таблица: что за данные используются, какого объёма, где лежит логика в коде."
    )
    _add_table(
        doc,
        [
            "БД / источник",
            "Назначение",
            "Состав / классы",
            "Объём",
            "Где в проекте",
        ],
        [
            [
                "MRL Eye Dataset",
                "Обучение CNN для класса «глаз открыт / закрыт»",
                "2 класса: open, closed",
                "Всего 84 898 изображений; train 67 919, val 16 979. "
                "Распределение по классам (пример): open ≈25,6 %, closed ≈74,4 % (сильный дисбаланс).",
                "Подготовка и обучение: src/training/, скрипты prepare* и train_eye_classifier.py; веса → models/",
            ],
            [
                "MediaPipe Face Landmarker (*.task)",
                "Инференс «моргания» по blendshapes без обучения на MRL",
                "Скаляры eyeBlink_L, eyeBlink_R (и др.)",
                "Файл модели ≈3,6 МБ",
                "src/analysis/eyes.py; основной режим в приложении",
            ],
        ],
    )
    doc.add_paragraph(
        "Примечание: в приложении для отбора кадров в реальном времени используется прежде всего MediaPipe "
        "(лёгкая модель и пороги по blendshapes). CNN (ResNet-18 и т.п.) обучаются на MRL и могут подключаться "
        "как дополнительная проверка в настройках."
    )

    doc.add_heading("2. Сравнение моделей ResNet-18 и MobileNetV3 (эпоха 3)", level=1)
    _add_table(
        doc,
        ["Модель", "Epoch", "Train acc", "Val acc", "Train loss", "Val loss"],
        [
            ["resnet18", "3", "0.978", "0.952", "0.075649", "0.114280"],
            ["mobilenetv3", "3", "0.964", "0.728", "0.095941", "0.597421"],
        ],
    )
    doc.add_paragraph("Мини-выводы по таблице:")
    for line in (
        "— ResNet-18 показывает высокую валидационную точность (≈95 %) и существенно меньшую валидационную "
        "ошибку, чем MobileNetV3 при сопоставимой тренировочной точности. Для проекта разумно использовать "
        "ResNet-18 как основную CNN при необходимости резервной классификации кропов глаз.",
        "— У MobileNetV3 наблюдается большой разрыв train/val (переобучение или слабая генерализация): val acc "
        "низкая, val loss высокий. Для production-модели без доработки (аугментации, регуляризации, проверки "
        "валидационного контура) MobileNetV3 менее надёжен.",
        "— Три эпохи достаточно, чтобы ResNet-18 вышел на стабильный уровень; дальнейшее обучение имеет смысл "
        "контролировать по валидации, чтобы не усугубить переобучение.",
    ):
        doc.add_paragraph(line)

    if (assets / "eye_models_summary.png").is_file():
        doc.add_paragraph()
        _add_picture_safe(
            doc,
            assets / "eye_models_summary.png",
            "Рис. Сводное сравнение метрик моделей (если подготовлен график eye_models_summary.png).",
        )

    doc.add_heading("3. Мини-выводы к графикам обучения", level=1)

    doc.add_paragraph("3.1. ResNet-18 — точность по эпохам", style="Heading 3")
    _add_picture_safe(doc, assets / "resnet18_accuracy.png", "Рис. ResNet-18: accuracy на train и validation.")
    doc.add_paragraph(
        "Вывод: тренировочная и валидационная точность растут согласованно; без аномального «обвала» val. "
        "Модель обобщает на отложенную выборку, что согласуется с высоким Val acc в таблице."
    )

    doc.add_paragraph("3.2. ResNet-18 — функция потерь (CrossEntropy)", style="Heading 3")
    _add_picture_safe(doc, assets / "resnet18_loss.png", "Рис. ResNet-18: loss на train и validation.")
    doc.add_paragraph(
        "Вывод: обе кривые снижаются; к поздним эпохам зазор train/val может немного увеличиваться — признак "
        "начала переобучения; разумно остановить обучение, ориентируясь на минимум val loss."
    )

    doc.add_paragraph("3.3. MobileNetV3 — точность по эпохам", style="Heading 3")
    _add_picture_safe(doc, assets / "mobilenetv3_accuracy.png", "Рис. MobileNetV3: accuracy на train и validation.")
    doc.add_paragraph(
        "Вывод: train accuracy растёт, а val остаётся на низком уровне (по графику — фактически без роста). "
        "Это указывает на плохую генерализацию; дополнительно стоит проверить корректность валидационного "
        "цикла и разметки, равномерность аугментаций и баланс классов на val."
    )

    doc.add_paragraph("3.4. MobileNetV3 — функция потерь", style="Heading 3")
    _add_picture_safe(doc, assets / "mobilenetv3_loss.png", "Рис. MobileNetV3: loss на train и validation.")
    doc.add_paragraph(
        "Вывод: train loss заметно падает, val loss остаётся высоким — классическая картина переобучения или "
        "несоответствия распределений train/val; для улучшения — сильнее регуляризовать сеть или упростить "
        "архитектуру под объём данных, пересмотреть сплит."
    )

    doc.add_heading("4. Кратко о реализации в коде (пример цепочки)", level=1)
    doc.add_paragraph(
        "1) Пайплайн анализа одного файла: класс PhotoAnalyzer в src/analysis/pipeline.py последовательно "
        "вызывает проверку глаз, затем экспозицию и резкость."
    )
    doc.add_paragraph(
        "2) Глаза: EyeStateAnalyzer в src/analysis/eyes.py. Основной режим — MediaPipe Face Landmarker с "
        "выходом blendshapes; по порогам eyeBlink_L/R принимается решение о закрытых веках. Опционально "
        "(настройка use_cnn_eye_check и файл весов) подключается классификатор из src/models/eye_classifier.py "
        "(ResNet-18 на два класса)."
    )
    doc.add_paragraph(
        "3) Параметры приложения и путь к весам: src/config/settings.py, сохранение в app_settings.json."
    )
    doc.add_paragraph(
        "Фрагмент логики (смысл, не дословная копия): после чтения кадра RGB вызывается landmarker.process; "
        "для крупнейшего лица берутся blendshapes; если оба показателя моргания выше порога и ниже порога "
        "«хотя бы один глаз слегка открыт» — кадр отбраковывается как «закрытые глаза»."
    )

    doc.add_heading("5. Заключение", level=1)
    doc.add_paragraph(
        "Для автоматической сортировки в приложении оправдан основной сценарий на MediaPipe + пороги "
        "blendshapes; обучение на MRL подтверждает, что ResNet-18 даёт лучший баланс качества на валидации, "
        "чем обученный в том же режиме MobileNetV3. Графики качества и потерь согласуются с численной таблицей "
        "и могут использоваться в пояснительной записке как иллюстрация эксперимента."
    )

    doc.save(out)
    print(f"Записано: {out}")


if __name__ == "__main__":
    main()
